from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

os.chdir(r"c:\Users\grace\Environmental Science Infographic Poster")

# ===== DOCUMENT 1: Full text with references =====
doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Title
p = doc.add_heading('Fast Fashion: Environmental Impact in Canada', level=1)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Understanding the Hidden Cost of Cheap Clothing')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

# Topic Statement
doc.add_heading('Topic Statement', level=2)
p = doc.add_paragraph(
    'Fast fashion refers to the rapid production of inexpensive clothing that mimics current luxury fashion trends, '
    'designed for short-term use and quick disposal (Environment and Climate Change Canada, 2023). Canada is one of '
    'the highest per-capita consumers of textiles in the world. Canadians purchase an estimated 70 new garments per '
    'person annually, yet the average garment is worn only 7\u201310 times before being discarded (Statistics Canada, '
    '2023; Ellen MacArthur Foundation, 2017). The environmental consequences of this linear \u201ctake-make-dispose\u201d '
    'model are severe, contributing to greenhouse gas emissions, water pollution, microplastic contamination, and '
    'massive textile waste in Canadian landfills (UNEP, 2023). The fashion industry alone accounts for approximately '
    '8\u201310% of global carbon emissions \u2014 more than international flights and maritime shipping combined (IPCC, 2022).'
)

# Key Statistics
doc.add_heading('Key Statistics', level=2)
stats = [
    'Canadians send approximately 12 kilograms of textile waste per person per year to landfill (Environment and Climate Change Canada, 2023).',
    'It takes roughly 2,700 litres of water to produce a single cotton T-shirt (Ellen MacArthur Foundation, 2017).',
    'The global fashion industry emits approximately 1.2 gigatonnes of CO\u2082 per year (UNEP, 2023).',
    'Less than 1% of clothing is recycled into new garments worldwide (Ellen MacArthur Foundation, 2017).',
]
for s in stats:
    doc.add_paragraph(s, style='List Bullet')

# Environmental Impacts
doc.add_heading('Environmental Impacts in Canada', level=2)

doc.add_heading('Carbon Emissions', level=3)
doc.add_paragraph(
    'Canada\u2019s textile and clothing sector contributes to greenhouse gas emissions throughout the supply chain. '
    'Globally, fashion produces approximately 1.2 gigatonnes of CO\u2082 annually \u2014 more than aviation and shipping '
    'combined. Due to the reliance on synthetic fibres derived from fossil fuels, emissions occur during both production '
    'and incineration of discarded garments (UNEP, 2023; IPCC, 2022).'
)

doc.add_heading('Water Consumption and Pollution', level=3)
doc.add_paragraph(
    'The fashion industry consumes approximately 93 billion cubic metres of water per year globally. Textile dyeing '
    'is the second-largest polluter of water worldwide. In Canada, contaminated wastewater from textile processes '
    'introduces heavy metals and toxic chemicals like formaldehyde and chlorine bleach into freshwater ecosystems '
    '(Ellen MacArthur Foundation, 2017; UNEP, 2023).'
)

doc.add_heading('Landfill and Waste Crisis', level=3)
doc.add_paragraph(
    'Canadians send an estimated 12 kilograms of textile waste per person to landfills each year, totalling over '
    '500,000 tonnes nationally. Synthetic textiles can take 20\u2013200 years to decompose and release methane, a '
    'potent greenhouse gas, during decomposition. Only about 15% of post-consumer textiles are diverted from landfills '
    'in Canada (Environment and Climate Change Canada, 2023; Statistics Canada, 2023).'
)

doc.add_heading('Biodiversity and Land Use', level=3)
doc.add_paragraph(
    'Cotton cultivation, a key raw material in fashion, uses 6% of the world\u2019s pesticides and 16% of insecticides. '
    'This intensive monoculture farming leads to soil degradation and biodiversity loss. The expansion of land for fibre '
    'crops has contributed to deforestation and habitat destruction in regions supplying Canadian retailers (UNEP, 2023; IPCC, 2022).'
)

doc.add_heading('Microplastic Pollution in Canadian Waters', level=3)
doc.add_paragraph(
    'Approximately 35% of primary microplastics in the world\u2019s oceans originate from washing synthetic textiles such '
    'as polyester and nylon (UNEP, 2023). In Canada, microplastics have been detected in the Great Lakes, the St. Lawrence '
    'River, and Arctic waters, with textile fibres being the most prevalent type identified (Environment and Climate Change '
    'Canada, 2023). A single domestic wash cycle can release up to 700,000 microplastic fibres into waterways (Ellen MacArthur Foundation, 2017).'
)

# Solutions
doc.add_heading('Solutions: How You Can Help', level=2)

solutions = [
    ('Buy Less, Choose Well', 'Invest in quality, durable garments. Reducing purchases by 30% could cut emissions equivalent to taking 1 million cars off roads (Ellen MacArthur Foundation, 2017).'),
    ('Thrift and Swap', 'Shop second-hand or organize clothing swaps. Canada\u2019s resale market is growing rapidly, extending garment lifespans and reducing demand for new production (Statistics Canada, 2023).'),
    ('Recycle and Donate', 'Use municipal textile recycling programs. Many Canadian cities now offer curbside textile collection to divert clothing from landfills (Environment and Climate Change Canada, 2023).'),
    ('Wash Mindfully', 'Use cold water, full loads, and microfibre-catching laundry bags to reduce microplastic release by up to 80% per wash cycle (Ellen MacArthur Foundation, 2017).'),
    ('Advocate for Policy', 'Support extended producer responsibility (EPR) policies. Canada is developing a federal strategy on textile waste management under the Zero Plastic Waste agenda (Environment and Climate Change Canada, 2023).'),
    ('Choose Sustainable Fabrics', 'Opt for certified organic cotton, TENCEL, or recycled fibres. Look for certifications like GOTS, OEKO-TEX, and Fair Trade (UNEP, 2023).'),
]

for title, text in solutions:
    p = doc.add_paragraph()
    run = p.add_run(title + ': ')
    run.bold = True
    p.add_run(text)

# References
doc.add_heading('References', level=2)
refs = [
    'Environment and Climate Change Canada (2023) Reducing plastic and textile waste. Government of Canada. Available at: https://www.canada.ca/en/environment-climate-change.html (Accessed: 20 February 2026).',
    'Statistics Canada (2023) Household spending on clothing and textiles. Statistics Canada. Available at: https://www.statcan.gc.ca/en/start (Accessed: 20 February 2026).',
    'Ellen MacArthur Foundation (2017) A new textiles economy: Redesigning fashion\u2019s future. Ellen MacArthur Foundation. Available at: https://www.ellenmacarthurfoundation.org/ (Accessed: 20 February 2026).',
    'United Nations Environment Programme (2023) Sustainability and circularity in the textile value chain. UNEP. Available at: https://www.unep.org/ (Accessed: 20 February 2026).',
    'Intergovernmental Panel on Climate Change (2022) Climate Change 2022: Mitigation of Climate Change. Cambridge University Press. Available at: https://www.ipcc.ch/ (Accessed: 20 February 2026).',
]
for i, ref in enumerate(refs, 1):
    doc.add_paragraph(f'{i}. {ref}')

doc.save('Infographic_Text.docx')
print('Created: Infographic_Text.docx')


# ===== DOCUMENT 2: Text without references (for Turnitin if needed) =====
doc2 = Document()
style2 = doc2.styles['Normal']
font2 = style2.font
font2.name = 'Calibri'
font2.size = Pt(11)

doc2.add_heading('Fast Fashion: Environmental Impact in Canada', level=1)
p = doc2.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Understanding the Hidden Cost of Cheap Clothing')
run.font.size = Pt(12)

doc2.add_heading('Topic Statement', level=2)
doc2.add_paragraph(
    'Fast fashion refers to the rapid production of inexpensive clothing that mimics current luxury fashion trends, '
    'designed for short-term use and quick disposal (Environment and Climate Change Canada, 2023). Canada is one of '
    'the highest per-capita consumers of textiles in the world. Canadians purchase an estimated 70 new garments per '
    'person annually, yet the average garment is worn only 7\u201310 times before being discarded (Statistics Canada, '
    '2023; Ellen MacArthur Foundation, 2017). The environmental consequences of this linear \u201ctake-make-dispose\u201d '
    'model are severe, contributing to greenhouse gas emissions, water pollution, microplastic contamination, and '
    'massive textile waste in Canadian landfills (UNEP, 2023). The fashion industry alone accounts for approximately '
    '8\u201310% of global carbon emissions \u2014 more than international flights and maritime shipping combined (IPCC, 2022).'
)

doc2.add_heading('Key Statistics', level=2)
for s in stats:
    doc2.add_paragraph(s, style='List Bullet')

doc2.add_heading('Environmental Impacts in Canada', level=2)

doc2.add_heading('Carbon Emissions', level=3)
doc2.add_paragraph(
    'Canada\u2019s textile and clothing sector contributes to greenhouse gas emissions throughout the supply chain. '
    'Globally, fashion produces approximately 1.2 gigatonnes of CO\u2082 annually \u2014 more than aviation and shipping '
    'combined. Due to the reliance on synthetic fibres derived from fossil fuels, emissions occur during both production '
    'and incineration of discarded garments (UNEP, 2023; IPCC, 2022).'
)

doc2.add_heading('Water Consumption and Pollution', level=3)
doc2.add_paragraph(
    'The fashion industry consumes approximately 93 billion cubic metres of water per year globally. Textile dyeing '
    'is the second-largest polluter of water worldwide. In Canada, contaminated wastewater from textile processes '
    'introduces heavy metals and toxic chemicals like formaldehyde and chlorine bleach into freshwater ecosystems '
    '(Ellen MacArthur Foundation, 2017; UNEP, 2023).'
)

doc2.add_heading('Landfill and Waste Crisis', level=3)
doc2.add_paragraph(
    'Canadians send an estimated 12 kilograms of textile waste per person to landfills each year, totalling over '
    '500,000 tonnes nationally. Synthetic textiles can take 20\u2013200 years to decompose and release methane, a '
    'potent greenhouse gas, during decomposition. Only about 15% of post-consumer textiles are diverted from landfills '
    'in Canada (Environment and Climate Change Canada, 2023; Statistics Canada, 2023).'
)

doc2.add_heading('Biodiversity and Land Use', level=3)
doc2.add_paragraph(
    'Cotton cultivation, a key raw material in fashion, uses 6% of the world\u2019s pesticides and 16% of insecticides. '
    'This intensive monoculture farming leads to soil degradation and biodiversity loss. The expansion of land for fibre '
    'crops has contributed to deforestation and habitat destruction in regions supplying Canadian retailers (UNEP, 2023; IPCC, 2022).'
)

doc2.add_heading('Microplastic Pollution in Canadian Waters', level=3)
doc2.add_paragraph(
    'Approximately 35% of primary microplastics in the world\u2019s oceans originate from washing synthetic textiles such '
    'as polyester and nylon (UNEP, 2023). In Canada, microplastics have been detected in the Great Lakes, the St. Lawrence '
    'River, and Arctic waters, with textile fibres being the most prevalent type identified (Environment and Climate Change '
    'Canada, 2023). A single domestic wash cycle can release up to 700,000 microplastic fibres into waterways (Ellen MacArthur Foundation, 2017).'
)

doc2.add_heading('Solutions: How You Can Help', level=2)
for title, text in solutions:
    p = doc2.add_paragraph()
    run = p.add_run(title + ': ')
    run.bold = True
    p.add_run(text)

doc2.save('Infographic_Text_NoRefs.docx')
print('Created: Infographic_Text_NoRefs.docx')


# ===== DOCUMENT 3: Reference list only =====
doc3 = Document()
style3 = doc3.styles['Normal']
font3 = style3.font
font3.name = 'Calibri'
font3.size = Pt(11)

doc3.add_heading('Reference List', level=1)
p = doc3.add_paragraph()
run = p.add_run('Fast Fashion: Environmental Impact in Canada')
run.font.size = Pt(12)
run.bold = True

doc3.add_paragraph('')

for i, ref in enumerate(refs, 1):
    doc3.add_paragraph(f'{i}. {ref}')

doc3.save('Reference_List.docx')
print('Created: Reference_List.docx')
